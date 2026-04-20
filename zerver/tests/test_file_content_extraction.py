from io import BytesIO

from docx import Document
from PIL import Image, ImageDraw

from zerver.lib.test_classes import UploadSerializeMixin, ZulipTestCase
from zerver.lib.upload import upload_message_attachment
from zerver.models import Attachment, AttachmentContent
from zerver.worker.file_content_extraction import (
    FileContentExtractionWorker,
    extract_from_docx,
    extract_from_pdf,
)


class FileContentExtractionWorkerTest(UploadSerializeMixin, ZulipTestCase):
    def test_consume_processes_attachment(self) -> None:
        """Sanity check: worker fetches attachment and runs without error."""
        user_profile = self.example_user("hamlet")
        url, _ = upload_message_attachment(
            "dummy.txt", "text/plain", b"zulip!", user_profile
        )
        path_id = url.replace("/user_uploads/", "")
        attachment = Attachment.objects.get(path_id=path_id)

        worker = FileContentExtractionWorker()
        worker.consume({"id": attachment.id})

        # Worker completed without raising; attachment unchanged
        attachment.refresh_from_db()
        self.assertEqual(attachment.path_id, path_id)
        self.assertIn("text/plain", attachment.content_type or "")

    def test_extract_docx(self) -> None:
        sample_document = Document()
        p1 = "Good morning everyone!"
        p2 = "Good night everyone."
        sample_document.add_paragraph(p1)
        sample_document.add_paragraph(p2)
        bio = BytesIO()
        sample_document.save(bio)
        file_bytes = bio.getvalue()
        self.assertEqual(extract_from_docx(file_bytes), p1 + "\n" + p2 + "\n")

    def test_extract_pdf(self) -> None:
        import pymupdf

        doc = pymupdf.open()
        page = doc.new_page()
        page.insert_text((72, 72), "Hello from PDF")
        page.insert_text((72, 100), "Second line of text")
        file_bytes = doc.tobytes()
        doc.close()

        extracted = extract_from_pdf(file_bytes)
        self.assertIn("Hello from PDF", extracted)
        self.assertIn("Second line of text", extracted)

    def test_extract_text_from_jpeg_image(self) -> None:
        user_profile = self.example_user("hamlet")

        # Keep generated text large/high-contrast to reduce OCR flakiness.
        image = Image.new("RGB", (900, 300), color="white")
        draw = ImageDraw.Draw(image)
        expected_text = "HELLO OCR 123"
        draw.text((40, 100), expected_text, fill="black")

        image_bytes = BytesIO()
        image.save(image_bytes, format="JPEG", quality=95)

        url, _ = upload_message_attachment(
            "ocr-test.jpg",
            "image/jpeg",
            image_bytes.getvalue(),
            user_profile,
        )
        path_id = url.replace("/user_uploads/", "")
        attachment = Attachment.objects.get(path_id=path_id)

        worker = FileContentExtractionWorker()
        worker.consume({"id": attachment.id})

        content = AttachmentContent.objects.get(attachment_id=attachment.id)
        self.assertEqual(content.extraction_status, 2)  # SUCCESS
        self.assertIsNotNone(content.extracted_text)
        normalized = " ".join((content.extracted_text or "").upper().split())
        self.assertIn("HELLO", normalized)
        self.assertIn("OCR", normalized)
